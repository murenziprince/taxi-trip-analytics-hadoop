"""
build_report.py
Generates Taxi_Analytics_Report.docx using YOUR actual cleaning audit,
performance comparison, and MapReduce output numbers - plus placeholder
sections marked [PASTE SCREENSHOT HERE: ...] for you to drop images into
in Word afterward.
Run this AFTER make_visualizations_windows.py.
"""
import json, os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

PROJECT = r"C:\taxi_project_assignment1"
FO = os.path.join(PROJECT, "final_outputs")
RA = os.path.join(PROJECT, "report_assets")
os.makedirs(RA, exist_ok=True)

with open(os.path.join(RA, "cleaning_audit.json")) as f:
    cleaning = json.load(f)
with open(os.path.join(RA, "performance_comparison.json")) as f:
    perf = json.load(f)

hourly = pd.read_csv(os.path.join(FO, "hourly.tsv"), sep="\t", header=None, names=["hour", "trips"])
daily = pd.read_csv(os.path.join(FO, "daily.tsv"), sep="\t", header=None, names=["day", "trips"])
loc = pd.read_csv(os.path.join(FO, "locations.tsv"), sep="\t", header=None, names=["zone", "trips"])
pay = pd.read_csv(os.path.join(FO, "payment.tsv"), sep="\t", header=None, names=["code", "stats"])
pay[["label", "ptrips", "prevenue", "pavgfare", "pavgtip"]] = pay["stats"].str.split(",", expand=True)
dist = pd.read_csv(os.path.join(FO, "distance.tsv"), sep="\t", header=None, names=["category", "stats"])
dist[["dtrips", "davgfare", "davgtotal", "davgdist"]] = dist["stats"].str.split(",", expand=True)
routes = pd.read_csv(os.path.join(FO, "routes.tsv"), sep="\t", header=None, names=["route", "stats"])
routes[["rtrips", "rrevenue"]] = routes["stats"].str.split(",", expand=True)
routes["rtrips"] = routes["rtrips"].astype(int); routes["rrevenue"] = routes["rrevenue"].astype(float)
anomalies = pd.read_csv(os.path.join(FO, "anomalies.tsv"), sep="\t", header=None, names=["type", "count"])
top10rev = pd.read_csv(os.path.join(FO, "revenue_top10.tsv"), sep="\t", header=None, names=["rank", "zone", "revenue"])

busiest_hour = hourly.loc[hourly["trips"].idxmax()]
quietest_hour = hourly.loc[hourly["trips"].idxmin()]
busiest_day = daily.loc[daily["trips"].idxmax()]
top_zone = loc.sort_values("trips", ascending=False).iloc[0]
top_payment = pay.sort_values("prevenue", key=lambda s: s.astype(float), ascending=False).iloc[0]
best_distance_cat = dist.loc[dist["davgtotal"].astype(float).idxmax()]
top_route_count = routes.sort_values("rtrips", ascending=False).iloc[0]
top_route_rev = routes.sort_values("rrevenue", ascending=False).iloc[0]
total_anomaly_records = anomalies["count"].sum()
flagged = total_anomaly_records - anomalies.loc[anomalies["type"] == "normal", "count"].values[0]
pct_flagged = 100 * flagged / total_anomaly_records

# ---------------- Build document ----------------
doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"; style.font.size = Pt(11)

def h1(text):
    doc.add_heading(text, level=1)

def h2(text):
    doc.add_heading(text, level=2)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = bold; r.italic = italic
    return p

def screenshot_placeholder(label):
    p = doc.add_paragraph()
    r = p.add_run(f"[ PASTE SCREENSHOT HERE: {label} ]")
    r.bold = True; r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    border_p = doc.add_paragraph("—" * 40)
    border_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    doc.add_paragraph()

def add_image(path, width=5.5):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
    else:
        para(f"[Chart not found: {path} - run make_visualizations_windows.py first]")

# Title page
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Distributed Taxi Trip Analytics")
r.bold = True; r.font.size = Pt(28)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Using Apache Hadoop, HDFS and Python MapReduce").font.size = Pt(16)
p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Big Data Essentials — Individual Practical Case Study")
r2.italic = True
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.add_run(f"Dataset: NYC TLC Yellow Taxi Trip Records — {', '.join(cleaning.get('months_included', ['January 2026']))}")
doc.add_page_break()

# 1. Introduction
h1("1. Introduction")
para(f"This report presents a distributed analytics solution for NYC TLC Yellow Taxi trip records "
     f"({cleaning['total_raw_records']:,} raw records across {len(cleaning.get('months_included', []))} month(s)), "
     f"built on HDFS for storage and Python Hadoop Streaming MapReduce for processing.")

# 2. Business Problem
h1("2. Business Problem")
para("A transportation analytics company wants to understand taxi demand, revenue, routes, payment "
     "behavior, trip characteristics, and data-quality anomalies at scale, using HDFS for distributed "
     "storage and MapReduce for horizontally scalable processing.")

# 3. Dataset Description
h1("3. Dataset Description")
add_table(["Property", "Value"], [
    ["Raw record count", f"{cleaning['total_raw_records']:,}"],
    ["Months included", ", ".join(cleaning.get("months_included", []))],
    ["Source format", "Apache Parquet"],
    ["Streaming input format", "CSV"],
])

# 4. Hadoop Environment
h1("4. Hadoop Environment")
para("Native Hadoop installation on Windows (single-node pseudo-distributed), with HDFS (NameNode + "
     "DataNode) and YARN (ResourceManager + NodeManager) daemons. Verified via `jps` and "
     "`hdfs dfsadmin -report` before job submission.")
screenshot_placeholder("jps output showing NameNode, DataNode, ResourceManager, NodeManager")
screenshot_placeholder("hdfs dfsadmin -report output")

# 5. HDFS Design
h1("5. HDFS Design")
para("HDFS directory structure under /taxi_project_assignment1/, separating raw input, cleaned input, "
     "per-analysis output, and archive.")
screenshot_placeholder("hdfs dfs -ls -R /taxi_project_assignment1 (full directory tree)")
screenshot_placeholder("hdfs dfs -ls -h on input/raw and input/cleaned, plus hdfs dfs -du -h")

# 6. Data Cleaning
h1("6. Data Cleaning")
para("Cleaning followed a minimum-necessary-deletion principle: fields that were simply unreported "
     "were imputed rather than dropping the whole record; records were only removed when a reported "
     "value was logically impossible.")
add_table(["Rule", "Description", "Records Affected", "% of Dataset"],
          [[r["rule"], r["description"], f"{r['records_affected']:,}", f"{r['pct_of_dataset']}%"] for r in cleaning["rules"]])
para(f"Result: {cleaning['total_raw_records']:,} raw records -> {cleaning['total_clean_records']:,} "
     f"clean records ({cleaning['total_removed']:,} removed, {cleaning['pct_removed']}%).", bold=True)

# 7. MapReduce Design
h1("7. MapReduce Design")
para("Every Mapper emits key\\tvalue text to stdout; Hadoop's Shuffle-and-Sort groups values by key "
     "for the Reducer, which tracks a single running accumulator per key (O(1) memory per key).")
add_table(["Analysis", "Mapper Key", "Reducer Output"], [
    ["a) Hourly demand", "pickup_hour", "trips per hour"],
    ["b) Daily demand", "day_of_week", "trips per day + weekday/weekend totals"],
    ["c) Pickup location", "PULocationID", "trips per zone, top10/bottom10"],
    ["d) Revenue by location", "PULocationID", "count, total fare/tip/revenue, avg fare/dist"],
    ["e) Payment method", "payment_type", "trips, revenue, avg fare/tip"],
    ["f) Distance-based fare", "distance bucket", "avg fare/total/distance"],
    ["g) Routes", "PU->DO", "count & revenue, top20 each"],
    ["h) Trip duration", "duration bucket", "avg fare/dist/tip"],
    ["i) Anomaly detection", "anomaly_type", "count per category, % flagged"],
])

# 8. Mapper/Reducer Implementation
h1("8. Mapper and Reducer Implementation")
para("All 9 analyses plus the 2-stage revenue job are implemented as standalone Python scripts in "
     "mappers/ and reducers/ (submitted alongside this report). Each uses only the Python standard "
     "library (sys, csv) - no dependencies needed on cluster worker nodes.")
screenshot_placeholder("Console output of a successful hadoop jar hadoop-streaming run (e.g. hourly job)")

# 9. Analytical Results
doc.add_page_break()
h1("9. Analytical Results")

h2("9.1 Hourly Taxi Demand")
add_image(os.path.join(RA, "01_trips_by_hour.png"))
para(f"Busiest hour: {int(busiest_hour['hour'])}:00 ({int(busiest_hour['trips']):,} trips). "
     f"Quietest hour: {int(quietest_hour['hour'])}:00 ({int(quietest_hour['trips']):,} trips).")

h2("9.2 Daily Demand")
add_image(os.path.join(RA, "02_trips_by_day.png"))
para(f"Busiest day: {busiest_day['day']} ({int(busiest_day['trips']):,} trips).")

h2("9.3 Pickup Location Analysis")
add_image(os.path.join(RA, "03_top10_pickup_zones.png"))
para(f"Top pickup zone: LocationID {int(top_zone['zone'])} ({int(top_zone['trips']):,} trips).")
screenshot_placeholder("hdfs dfs -cat .../output/locations/part-00000")

h2("9.4 Revenue by Pickup Location")
add_table(["Rank", "Zone", "Total Revenue ($)"],
          [[r["rank"], r["zone"], f"{r['revenue']:,.2f}"] for _, r in top10rev.iterrows()])
screenshot_placeholder("hdfs dfs -cat .../output/revenue/part-00000 (full zone stats)")

h2("9.5 Payment Method Analysis")
add_image(os.path.join(RA, "04_revenue_by_payment.png"))
para(f"Top payment method by revenue: {top_payment['label']} (${float(top_payment['prevenue']):,.2f}).")

h2("9.6 Distance-Based Fare Analysis")
add_image(os.path.join(RA, "05_trips_by_distance_category.png"))
add_image(os.path.join(RA, "07_revenue_vs_distance.png"))
para(f"Highest average fare distance category: {best_distance_cat['category']} "
     f"(avg total ${float(best_distance_cat['davgtotal']):.2f}).")

h2("9.7 Busiest Pickup-Dropoff Routes")
add_image(os.path.join(RA, "06_top10_routes.png"))
para(f"Most frequent route: {top_route_count['route']} ({int(top_route_count['rtrips']):,} trips). "
     f"Most profitable route: {top_route_rev['route']} (${top_route_rev['rrevenue']:,.2f}).")

h2("9.8 Trip Duration Analysis")
screenshot_placeholder("hdfs dfs -cat .../output/duration/part-00000")

h2("9.9 Anomaly Detection")
para(f"{int(total_anomaly_records):,} clean records scanned; {int(flagged):,} flagged ({pct_flagged:.2f}%).")
add_table(["Anomaly Type", "Count"], [[r["type"], f"{int(r['count']):,}"] for _, r in anomalies.iterrows()])

# 10. Multi-Stage MapReduce
doc.add_page_break()
h1("10. Multi-Stage MapReduce")
para("Stage 1 (revenue by pickup zone) writes to HDFS output/revenue/. Stage 2 reads that HDFS output "
     "directly, re-keys every zone to a single constant key so all zone totals funnel to one reducer, "
     "and emits the globally sorted Top 10.")
screenshot_placeholder("Stage 1 intermediate output: hdfs dfs -cat .../output/revenue/part-00000 (BEFORE running Stage 2)")
add_table(["Rank", "Zone", "Total Revenue ($)"],
          [[r["rank"], r["zone"], f"{r['revenue']:,.2f}"] for _, r in top10rev.iterrows()])

# 11. YARN Analysis
h1("11. YARN Analysis")
para("Each hadoop jar submission is scheduled by YARN's ResourceManager and tracked at "
     "http://localhost:8088/cluster, where Application ID, state, containers, timestamps and final "
     "status were captured as evidence.")
screenshot_placeholder("YARN ResourceManager application list (http://localhost:8088/cluster)")
screenshot_placeholder("YARN application detail page for one job (Application ID, containers, status)")

# 12. Performance Comparison
h1("12. Performance Comparison")
add_table(["Metric", "Python/Pandas", "Streaming Pipeline"], [
    ["Dataset size", f"{perf['dataset_size_mb']} MB", f"{perf['dataset_size_mb']} MB"],
    ["Number of records", f"{perf['num_records']:,}", f"{perf['num_records']:,}"],
    ["Execution time", f"{perf['pandas_time_sec']}s", f"{perf['streaming_time_sec']}s"],
    ["Output size", "in-memory DataFrame", f"{perf['streaming_output_kb']} KB"],
])
screenshot_placeholder("Console output of performance_comparison.py")
para("Discussion: at this dataset scale on a single machine, Pandas is faster in wall-clock time "
     "because it has no process-spawn/text-parsing overhead. Hadoop MapReduce's value is horizontal "
     "scalability and fault tolerance (as observed first-hand during job execution, where failed/killed "
     "map and reduce tasks were automatically retried by YARN until the job completed successfully) "
     "- advantages that compound at multi-GB-to-TB scale, not at this single-month/single-node test size.")

# 13. Business Insights
doc.add_page_break()
h1("13. Business Insights (Final Business Questions)")
qas = [
    ("a) Busiest hour?", f"{int(busiest_hour['hour'])}:00, {int(busiest_hour['trips']):,} trips."),
    ("b) Busiest day?", f"{busiest_day['day']}, {int(busiest_day['trips']):,} trips."),
    ("c) Zones with most trips?", f"LocationID {int(top_zone['zone'])} leads with {int(top_zone['trips']):,} trips."),
    ("d) Zones with most revenue?", f"See Top-10 revenue table in Section 10."),
    ("e) Payment method with most revenue?", f"{top_payment['label']}, ${float(top_payment['prevenue']):,.2f}."),
    ("f) Do credit-card users tip more?", "Compare avg_tip across payment.tsv - cash tips are not captured by TLC data at all, a known data-source limitation."),
    ("g) Distance category with highest avg fare?", f"{best_distance_cat['category']}."),
    ("h) Most frequent routes?", f"{top_route_count['route']}, {int(top_route_count['rtrips']):,} trips."),
    ("i) Are frequent routes also profitable?", f"Compare '{top_route_count['route']}' (frequency leader) against '{top_route_rev['route']}' (revenue leader) - see Section 9.7."),
    ("j) % of records with anomalies?", f"{pct_flagged:.2f}%."),
    ("k) Management insights?", "Fleet allocation should peak around the busiest hour/day; high-revenue-but-lower-frequency zones deserve targeted commercial attention; short trips are the volume backbone but least profitable per trip."),
    ("l) When does MapReduce beat conventional processing?", "When data exceeds single-machine memory/CPU budget, when jobs run repeatedly at production scale, or when fault-tolerant distributed execution is required - not necessarily at this single-month test scale, where Pandas was faster."),
]
for q, a in qas:
    para(q, bold=True)
    para(a)

# 14. Limitations
h1("14. Limitations")
for t in [
    f"{len(cleaning.get('months_included', []))} month(s) of data used.",
    "MapReduce jobs run on a real single-node Windows Hadoop cluster; YARN evidence reflects that environment.",
    "TLC data does not record cash tips.",
    "LocationIDs shown without a joined zone-name lookup table.",
    "Anomaly thresholds are heuristic, not tuned against labeled fraud data.",
]:
    doc.add_paragraph(t, style="List Bullet")

# 15. Conclusion
h1("15. Conclusion")
para("This project implemented a complete Hadoop Streaming analytics pipeline for NYC TLC Yellow Taxi "
     "data on a real Hadoop cluster: documented data cleaning, nine single-stage MapReduce analyses, a "
     "two-stage MapReduce workflow, and a Pandas-vs-MapReduce performance comparison, all validated "
     "with real YARN job execution and HDFS output.")

# 16. References
h1("16. References")
for t in [
    "NYC Taxi & Limousine Commission. TLC Trip Record Data. https://www.nyc.gov/site/tlc/about/tlc-trip-record-data.page",
    "Apache Hadoop Documentation. https://hadoop.apache.org/docs/stable/",
    "Apache Hadoop Streaming Guide. https://hadoop.apache.org/docs/stable/hadoop-streaming/HadoopStreaming.html",
]:
    doc.add_paragraph(t, style="List Bullet")

out_path = os.path.join(RA, "Taxi_Analytics_Report.docx")
doc.save(out_path)
print(f"Report saved to: {out_path}")
print(f"Contains {len(qas)} business Q&As, all real numbers from your cluster run.")
print("Search the document for '[ PASTE SCREENSHOT HERE' to find every spot needing a screenshot.")
